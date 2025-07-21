"""
多格式报告生成器 - 支持txt, json, excel, html, word格式
"""

import os
import json
import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import openpyxl

from loguru import logger
from auditluma.config import Config

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl未安装，Excel报告功能不可用")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    logger.warning("python-docx未安装，Word报告功能不可用")


class MultiFormatReportGenerator:
    """多格式报告生成器"""
    
    def __init__(self):
        """初始化报告生成器"""
        self.output_dir = Path(Config.get_report_dir())
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def load_analysis_data(self, data_file_path: str) -> Dict[str, Any]:
        """从history文件加载分析数据
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            分析数据字典
        """
        try:
            with open(data_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data
        except Exception as e:
            logger.error(f"加载分析数据失败: {e}")
            return {}
    
    def generate_filename(self, format_type: str, analysis_data: Dict[str, Any]) -> str:
        """生成报告文件名
        
        Args:
            format_type: 报告格式类型
            analysis_data: 分析数据
            
        Returns:
            文件名
        """
        project_name = analysis_data.get("scan_info", {}).get("project_name", "Unknown")
        
        # 清理项目名称，移除不适合文件名的字符
        safe_project_name = "".join(c for c in project_name if c.isalnum() or c in "._-").rstrip()
        if not safe_project_name:
            safe_project_name = "Unknown"
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{safe_project_name}_security_report_{timestamp}.{format_type}"
    
    def generate_txt_report(self, data_file_path: str) -> str:
        """生成TXT格式报告
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            生成的报告文件路径
        """
        data = self.load_analysis_data(data_file_path)
        if not data:
            raise ValueError("无法加载分析数据")
        
        filename = self.generate_filename("txt", data)
        report_path = self.output_dir / filename
        
        with open(report_path, 'w', encoding='utf-8') as f:
            # 写入报告头部
            f.write("=" * 80 + "\n")
            f.write("AuditLuma 安全审计报告\n")
            f.write("=" * 80 + "\n\n")
            
            # 基本信息
            scan_info = data.get("scan_info", {})
            f.write(f"项目名称: {scan_info.get('project_name', 'Unknown')}\n")
            f.write(f"扫描时间: {data.get('analysis_time', 'Unknown')}\n")
            f.write(f"扫描文件数: {scan_info.get('scanned_files', 0)}\n")
            f.write(f"扫描代码行数: {scan_info.get('scanned_lines', 0)}\n")
            f.write(f"扫描耗时: {scan_info.get('scan_duration', 'Unknown')}\n")
            f.write(f"发现漏洞总数: {data.get('vulnerabilities_count', 0)}\n\n")
            
            # 漏洞统计
            f.write("-" * 40 + "\n")
            f.write("漏洞统计\n")
            f.write("-" * 40 + "\n")
            
            severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
            for vuln in data.get("vulnerabilities", []):
                severity = vuln.get("severity", "").lower()
                if severity in severity_counts:
                    severity_counts[severity] += 1
            
            f.write(f"严重漏洞: {severity_counts['critical']}\n")
            f.write(f"高危漏洞: {severity_counts['high']}\n")
            f.write(f"中危漏洞: {severity_counts['medium']}\n")
            f.write(f"低危漏洞: {severity_counts['low']}\n")
            f.write(f"信息提示: {severity_counts['info']}\n\n")
            
            # 漏洞详情
            if data.get("vulnerabilities"):
                f.write("-" * 40 + "\n")
                f.write("漏洞详情\n")
                f.write("-" * 40 + "\n\n")
                
                for i, vuln in enumerate(data.get("vulnerabilities", []), 1):
                    f.write(f"{i}. {vuln.get('vulnerability_type', 'Unknown')}\n")
                    f.write(f"   ID: {vuln.get('id', 'N/A')}\n")
                    f.write(f"   严重程度: {vuln.get('severity', 'Unknown')}\n")
                    
                    # CVSS 4.0信息
                    if vuln.get('cvss4_score'):
                        f.write(f"   CVSS 4.0分数: {vuln.get('cvss4_score')} ({vuln.get('cvss4_severity', 'Unknown')})\n")
                        if vuln.get('cvss4_vector'):
                            f.write(f"   CVSS向量: {vuln.get('cvss4_vector')}\n")
                    
                    f.write(f"   文件: {vuln.get('file_path', 'Unknown')}\n")
                    f.write(f"   行数: {vuln.get('start_line', 0)}-{vuln.get('end_line', 0)}\n")
                    f.write(f"   描述: {vuln.get('description', '无描述')}\n")
                    f.write("\n")
            
            # 修复建议
            full_results = data.get("full_analysis_results", {})
            remediation_data = full_results.get("remediation_data")
            if remediation_data and remediation_data.get("remediations"):
                f.write("-" * 40 + "\n")
                f.write("修复建议\n")
                f.write("-" * 40 + "\n\n")
                
                for i, remediation in enumerate(remediation_data.get("remediations", []), 1):
                    f.write(f"{i}. 漏洞ID: {remediation.get('vulnerability_id', 'N/A')}\n")
                    f.write(f"   修复优先级: {remediation.get('priority', 'Unknown')}\n")
                    f.write(f"   修复难度: {remediation.get('difficulty', 'Unknown')}\n")
                    f.write(f"   修复建议:\n")
                    f.write(f"   {remediation.get('specific_remediation', '无具体建议')}\n")
                    f.write("\n")
        
        logger.info(f"TXT报告已生成: {report_path}")
        return str(report_path)
    
    def generate_json_report(self, data_file_path: str) -> str:
        """生成JSON格式报告
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            生成的报告文件路径
        """
        data = self.load_analysis_data(data_file_path)
        if not data:
            raise ValueError("无法加载分析数据")
        
        filename = self.generate_filename("json", data)
        report_path = self.output_dir / filename
        
        # 准备结构化的JSON报告数据
        report_data = {
            "report_meta": {
                "format": "json",
                "version": "1.0",
                "generated_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "generator": "AuditLuma MultiFormat Generator"
            },
            "scan_info": data.get("scan_info", {}),
            "analysis_time": data.get("analysis_time"),
            "summary": {
                "total_vulnerabilities": data.get("vulnerabilities_count", 0),
                "severity_distribution": {}
            },
            "vulnerabilities": data.get("vulnerabilities", []),
            "dependency_info": data.get("dependency_info", {}),
            "remediation_info": data.get("remediation_info", {}),
            "remediations": data.get("full_analysis_results", {}).get("remediation_data", {}).get("remediations", [])
        }
        
        # 计算严重程度分布
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for vuln in data.get("vulnerabilities", []):
            severity = vuln.get("severity", "").lower()
            if severity in severity_counts:
                severity_counts[severity] += 1
        
        report_data["summary"]["severity_distribution"] = severity_counts
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"JSON报告已生成: {report_path}")
        return str(report_path)
    
    def generate_excel_report(self, data_file_path: str) -> str:
        """生成Excel格式报告
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            生成的报告文件路径
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl未安装，无法生成Excel报告。请运行: pip install openpyxl")
        
        data = self.load_analysis_data(data_file_path)
        if not data:
            raise ValueError("无法加载分析数据")
        
        filename = self.generate_filename("xlsx", data)
        report_path = self.output_dir / filename
        
        # 创建工作簿和工作表
        wb = openpyxl.Workbook()
        
        # 概览工作表
        ws_summary = wb.active
        ws_summary.title = "概览"
        
        # 设置标题样式
        title_font = Font(size=16, bold=True)
        header_font = Font(size=12, bold=True)
        
        # 写入概览信息
        ws_summary['A1'] = "AuditLuma 安全审计报告"
        ws_summary['A1'].font = title_font
        
        scan_info = data.get("scan_info", {})
        ws_summary['A3'] = "基本信息"
        ws_summary['A3'].font = header_font
        
        info_data = [
            ["项目名称", scan_info.get('project_name', 'Unknown')],
            ["扫描时间", data.get('analysis_time', 'Unknown')],
            ["扫描文件数", scan_info.get('scanned_files', 0)],
            ["扫描代码行数", scan_info.get('scanned_lines', 0)],
            ["扫描耗时", scan_info.get('scan_duration', 'Unknown')],
            ["发现漏洞总数", data.get('vulnerabilities_count', 0)]
        ]
        
        for i, (key, value) in enumerate(info_data, 4):
            ws_summary[f'A{i}'] = key
            ws_summary[f'B{i}'] = value
        
        # 严重程度统计
        ws_summary['A11'] = "漏洞统计"
        ws_summary['A11'].font = header_font
        
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for vuln in data.get("vulnerabilities", []):
            severity = vuln.get("severity", "").lower()
            if severity in severity_counts:
                severity_counts[severity] += 1
        
        severity_labels = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危", "info": "信息"}
        for i, (severity, count) in enumerate(severity_counts.items(), 12):
            ws_summary[f'A{i}'] = severity_labels.get(severity, severity)
            ws_summary[f'B{i}'] = count
        
        # 漏洞详情工作表
        if data.get("vulnerabilities"):
            ws_vulns = wb.create_sheet("漏洞详情")
            
            # 设置列标题
            headers = ["序号", "漏洞类型", "严重程度", "CVSS4.0分数", "CVSS4.0等级", "文件路径", "起始行", "结束行", "漏洞ID", "描述", "CVSS向量"]
            for i, header in enumerate(headers, 1):
                cell = ws_vulns.cell(row=1, column=i, value=header)
                cell.font = header_font
            
            # 写入漏洞数据
            for row_idx, vuln in enumerate(data.get("vulnerabilities", []), 2):
                ws_vulns.cell(row=row_idx, column=1, value=row_idx-1)
                ws_vulns.cell(row=row_idx, column=2, value=vuln.get('vulnerability_type', 'Unknown'))
                ws_vulns.cell(row=row_idx, column=3, value=vuln.get('severity', 'Unknown'))
                ws_vulns.cell(row=row_idx, column=4, value=vuln.get('cvss4_score', 'N/A'))
                ws_vulns.cell(row=row_idx, column=5, value=vuln.get('cvss4_severity', 'N/A'))
                ws_vulns.cell(row=row_idx, column=6, value=vuln.get('file_path', 'Unknown'))
                ws_vulns.cell(row=row_idx, column=7, value=vuln.get('start_line', 0))
                ws_vulns.cell(row=row_idx, column=8, value=vuln.get('end_line', 0))
                ws_vulns.cell(row=row_idx, column=9, value=vuln.get('id', 'N/A'))
                ws_vulns.cell(row=row_idx, column=10, value=vuln.get('description', '无描述'))
                ws_vulns.cell(row=row_idx, column=11, value=vuln.get('cvss4_vector', 'N/A'))
            
            # 调整列宽
            ws_vulns.column_dimensions['B'].width = 25  # 漏洞类型
            ws_vulns.column_dimensions['D'].width = 12  # CVSS4.0分数
            ws_vulns.column_dimensions['E'].width = 12  # CVSS4.0等级
            ws_vulns.column_dimensions['F'].width = 40  # 文件路径
            ws_vulns.column_dimensions['I'].width = 40  # 漏洞ID
            ws_vulns.column_dimensions['J'].width = 50  # 描述
            ws_vulns.column_dimensions['K'].width = 60  # CVSS向量
        
        # 修复建议工作表
        full_results = data.get("full_analysis_results", {})
        remediation_data = full_results.get("remediation_data")
        if remediation_data and remediation_data.get("remediations"):
            ws_remediation = wb.create_sheet("修复建议")
            
            # 设置列标题
            headers = ["序号", "漏洞ID", "修复优先级", "修复难度", "修复建议"]
            for i, header in enumerate(headers, 1):
                cell = ws_remediation.cell(row=1, column=i, value=header)
                cell.font = header_font
            
            # 写入修复建议数据
            for row_idx, remediation in enumerate(remediation_data.get("remediations", []), 2):
                ws_remediation.cell(row=row_idx, column=1, value=row_idx-1)
                ws_remediation.cell(row=row_idx, column=2, value=remediation.get('vulnerability_id', 'N/A'))
                ws_remediation.cell(row=row_idx, column=3, value=remediation.get('priority', 'Unknown'))
                ws_remediation.cell(row=row_idx, column=4, value=remediation.get('difficulty', 'Unknown'))
                ws_remediation.cell(row=row_idx, column=5, value=remediation.get('specific_remediation', '无具体建议'))
            
            # 调整列宽
            ws_remediation.column_dimensions['B'].width = 40
            ws_remediation.column_dimensions['E'].width = 60
        
        # 保存文件
        wb.save(report_path)
        
        logger.info(f"Excel报告已生成: {report_path}")
        return str(report_path)
    
    def generate_html_report(self, data_file_path: str) -> str:
        """生成HTML格式报告（使用现有的HTML报告模板）
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            生成的报告文件路径
        """
        from auditluma.visualizer.report_generator import ReportGenerator
        from auditluma.models.code import VulnerabilityResult
        
        data = self.load_analysis_data(data_file_path)
        if not data:
            raise ValueError("无法加载分析数据")
        
        # 重建VulnerabilityResult对象
        vulnerabilities = []
        for vuln_data in data.get("vulnerabilities", []):
            # 创建虚拟CodeUnit对象（用于报告生成）
            from auditluma.models.code import CodeUnit, SourceFile, FileType
            from pathlib import Path
            
            dummy_source = SourceFile(
                path=Path(vuln_data.get("file_path", "unknown.py")),
                relative_path=vuln_data.get("file_path", "unknown.py"),
                name=Path(vuln_data.get("file_path", "unknown.py")).name,
                extension=Path(vuln_data.get("file_path", "unknown.py")).suffix,
                file_type=FileType.PYTHON,
                size=100,
                content="# Source content not available",
                modified_time=0.0
            )
            
            dummy_code_unit = CodeUnit(
                id=f"unit_{vuln_data.get('id', 'unknown')}",
                name="vulnerability_location",
                type="function",
                source_file=dummy_source,
                start_line=vuln_data.get("start_line", 1),
                end_line=vuln_data.get("end_line", 1),
                content=vuln_data.get("snippet", "# Code not available")
            )
            
            vuln = VulnerabilityResult(
                id=vuln_data.get("id"),
                title=vuln_data.get("title", vuln_data.get("vulnerability_type", "Unknown Vulnerability")),
                description=vuln_data.get("description"),
                code_unit=dummy_code_unit,
                file_path=vuln_data.get("file_path"),
                start_line=vuln_data.get("start_line"),
                end_line=vuln_data.get("end_line"),
                vulnerability_type=vuln_data.get("vulnerability_type"),
                severity=vuln_data.get("severity"),
                snippet=vuln_data.get("snippet", "")
            )
            
            # 添加CVSS 4.0信息
            vuln.cvss4_score = vuln_data.get("cvss4_score")
            vuln.cvss4_vector = vuln_data.get("cvss4_vector") 
            vuln.cvss4_severity = vuln_data.get("cvss4_severity")
            
            # 添加其他字段
            vuln.cwe_id = vuln_data.get("cwe_id")
            vuln.owasp_category = vuln_data.get("owasp_category")
            vuln.confidence = vuln_data.get("confidence", 1.0)
            vuln.recommendation = vuln_data.get("recommendation", "")
            vuln.references = vuln_data.get("references", [])
            
            # 添加metadata属性
            if hasattr(vuln, 'metadata'):
                vuln.metadata = vuln_data.get("metadata", {})
            vulnerabilities.append(vuln)
        
        # 获取修复建议数据
        full_results = data.get("full_analysis_results", {})
        remediation_data = full_results.get("remediation_data")
        
        # 使用现有的报告生成器
        report_generator = ReportGenerator()
        report_path = report_generator.generate_report(
            vulnerabilities=vulnerabilities,
            dependency_graph=None,  # 历史数据中不保存图对象
            remediation_data=remediation_data,  # 包含修复建议
            scan_info=data.get("scan_info", {})
        )
        
        logger.info(f"HTML报告已生成: {report_path}")
        return str(report_path)
    
    def generate_word_report(self, data_file_path: str) -> str:
        """生成Word格式报告
        
        Args:
            data_file_path: 数据文件路径
            
        Returns:
            生成的报告文件路径
        """
        if not WORD_AVAILABLE:
            raise ImportError("python-docx未安装，无法生成Word报告。请运行: pip install python-docx")
        
        data = self.load_analysis_data(data_file_path)
        if not data:
            raise ValueError("无法加载分析数据")
        
        filename = self.generate_filename("docx", data)
        report_path = self.output_dir / filename
        
        # 创建Word文档
        doc = Document()
        
        # 标题
        title = doc.add_heading('AuditLuma 安全审计报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_heading('基本信息', level=1)
        scan_info = data.get("scan_info", {})
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("项目名称", scan_info.get('project_name', 'Unknown')),
            ("扫描时间", data.get('analysis_time', 'Unknown')),
            ("扫描文件数", str(scan_info.get('scanned_files', 0))),
            ("扫描代码行数", str(scan_info.get('scanned_lines', 0))),
            ("扫描耗时", scan_info.get('scan_duration', 'Unknown')),
            ("发现漏洞总数", str(data.get('vulnerabilities_count', 0)))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 漏洞统计
        doc.add_heading('漏洞统计', level=1)
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for vuln in data.get("vulnerabilities", []):
            severity = vuln.get("severity", "").lower()
            if severity in severity_counts:
                severity_counts[severity] += 1
        
        stats_table = doc.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        severity_labels = {"critical": "严重漏洞", "high": "高危漏洞", "medium": "中危漏洞", "low": "低危漏洞", "info": "信息提示"}
        for i, (severity, count) in enumerate(severity_counts.items()):
            stats_table.cell(i, 0).text = severity_labels.get(severity, severity)
            stats_table.cell(i, 1).text = str(count)
        
        # 漏洞详情
        if data.get("vulnerabilities"):
            doc.add_heading('漏洞详情', level=1)
            
            for i, vuln in enumerate(data.get("vulnerabilities", []), 1):
                doc.add_heading(f'{i}. {vuln.get("vulnerability_type", "Unknown")}', level=2)
                
                vuln_info = [
                    f"漏洞ID: {vuln.get('id', 'N/A')}",
                    f"严重程度: {vuln.get('severity', 'Unknown')}",
                    f"文件路径: {vuln.get('file_path', 'Unknown')}",
                    f"行数: {vuln.get('start_line', 0)}-{vuln.get('end_line', 0)}",
                    f"描述: {vuln.get('description', '无描述')}"
                ]
                
                for info in vuln_info:
                    doc.add_paragraph(info)
                
                if vuln.get('snippet'):
                    doc.add_paragraph("代码片段:", style='Heading 3')
                    code_para = doc.add_paragraph(vuln.get('snippet'))
                    code_para.style = 'Normal'
        
        # 修复建议
        full_results = data.get("full_analysis_results", {})
        remediation_data = full_results.get("remediation_data")
        if remediation_data and remediation_data.get("remediations"):
            doc.add_heading('修复建议', level=1)
            
            for i, remediation in enumerate(remediation_data.get("remediations", []), 1):
                doc.add_heading(f'{i}. 漏洞ID: {remediation.get("vulnerability_id", "N/A")}', level=2)
                
                remediation_info = [
                    f"修复优先级: {remediation.get('priority', 'Unknown')}",
                    f"修复难度: {remediation.get('difficulty', 'Unknown')}",
                    f"修复建议: {remediation.get('specific_remediation', '无具体建议')}"
                ]
                
                for info in remediation_info:
                    doc.add_paragraph(info)
        
        # 保存文档
        doc.save(report_path)
        
        logger.info(f"Word报告已生成: {report_path}")
        return str(report_path) 